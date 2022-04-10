# -*- coding: utf-8 -*-
"""
Created on Tue Mar 16 10:36:21 2021

@author: Arjun Krishnan
# Assignment - 2
# Due Date - 18-03-2021
"""

import docx
import pyexcel as pe
import os
import regex as re


# Function recieves a docx file and saves an excel file
# with words whose frequency greater than or equalt to 0.001
def analyze(docfile):
    file_name = re.match(r".+(?=\.docx)", docfile).group()
    current_directory = os.getcwd()
    xls_est = "xlsx"
    op_file_path = os.path.join(current_directory, file_name + "_word_stats"+"." + xls_est)
    doc = docx.Document(docfile)
    print("Document has %d paragraphs." % len(doc.paragraphs))
    word_dict = {}
    word_freq_dict = {}
    sorted_word_freq_dict = {}
    word_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            # From the paragraph, split the string based on space
            for word in run.text.split():
                # Removes any special character at start and end of a word
                m = re.sub(r"[^a-zA-Z0-9]+$",'',re.sub(r'^[^a-zA-Z0-9]', '', word))
                if len(m) > 0:
                    word_count += 1
                    if m.lower() in word_dict:
                        word_dict[m.lower()] += 1
                    else:
                        word_dict[m.lower()] = 1
        
    print("Word count - %d" %word_count)
    
    word_freq_dict = {k: v/word_count for k, v in word_dict.items()}
    sorted_word_freq_dict = {k: v for k, v in sorted(word_freq_dict.items(), key = lambda kv : kv[1], reverse=True) if v >= 0.001}
    data = [[k, v] for k, v in sorted_word_freq_dict.items()]
    contents = {"Word Frequency Stats":data}
    pe.save_book_as(bookdict = contents, dest_file_name = op_file_path)
    
if __name__ == "__main__":
    docfile = "ulysses.docx"
    analyze(docfile)


